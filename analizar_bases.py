# -*- coding: utf-8 -*-
"""
Analizador de bases de licitaciones — Medina Consultoría SpA.

Corre LOCAL. Lee las bases que descargaste (PDF/DOCX/ZIP) en la carpeta
de licitaciones, extrae su texto, detecta la estructura de secciones y
clasifica los "puntos clave" (plazos, garantías, presupuesto, criterios de
evaluación, tareas/alcance, entregables, multas, equipo). Genera:

  - bases_data.js         (los datos, para el visor)
  - Analizador de Bases.html   (copia del visor)

Ambos en la carpeta de licitaciones. Abre el HTML para leer, recorrer con
índice hipervinculado y destacar tareas (se guardan en tu navegador).

Requisitos:  pip install pymupdf python-docx
"""
import sys, os, re, json, zipfile, pathlib, datetime, unicodedata

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

BASE_DIR = pathlib.Path(os.environ.get(
    "RADAR_DESTINO", pathlib.Path.home() / "Análisis RMG" / "Licitaciones"))
AQUI = pathlib.Path(__file__).parent

# --- OCR (Tesseract vía PyMuPDF) para PDFs escaneados ---
def _find_tessdata():
    for p in [os.environ.get("TESSDATA_PREFIX"),
              r"C:\Program Files\Tesseract-OCR\tessdata",
              r"C:\Program Files (x86)\Tesseract-OCR\tessdata",
              str(pathlib.Path.home() / "AppData/Local/Programs/Tesseract-OCR/tessdata")]:
        if p and os.path.isdir(p):
            return p
    return None
TESSDATA = _find_tessdata()
if TESSDATA:
    os.environ["TESSDATA_PREFIX"] = TESSDATA
OCR_OK = TESSDATA is not None

# Caché en Markdown: cada base leída (o OCR-eada) se guarda como .md y en las
# corridas siguientes se lee de ahí (rápido, sin re-procesar). --reprocesar lo ignora.
MD_DIR = BASE_DIR / "_markdown"
REPROC = "--reprocesar" in sys.argv

# --- clasificación del tipo de documento por nombre de archivo ---
def tipo_doc(nombre):
    n = sin_tildes(nombre).lower()
    if "administrativ" in n: return ("Administrativa", "adm")
    if re.search(r"\bbt[_ ]|tecnic|terminos.*referenc|\btdr\b|especificac", n): return ("Técnica", "tec")
    if "anexo" in n: return ("Anexos", "anx")
    if "presupuesto" in n: return ("Presupuesto", "eco")
    if re.search(r"decreto|resoluc|aprueba|\bres\b", n): return ("Resolución", "res")
    if "concurso" in n or "bases" in n: return ("Bases", "bas")
    return ("Documento", "doc")

# --- categorías de puntos clave (se buscan en los títulos de sección) ---
CATEGORIAS = [
    ("Plazos y fechas",        r"plazo|calendario|cronograma|vigencia|fecha"),
    ("Garantías",              r"garantia|boleta|fiel cumplimiento|seriedad"),
    ("Presupuesto y pagos",    r"presupuesto|monto|precio|\bpago|estados de pago|financiamiento|reajuste"),
    ("Criterios de evaluación",r"evaluacion|pauta|ponderaci|puntaje|criterio"),
    ("Requisitos y oferentes", r"requisito|participante|admisibilidad|oferente|union temporal|\butp\b"),
    ("Tareas y alcance (TDR)", r"tarea|alcance|objetivo|actividad|metodolog|desarrollo de proyecto|estudios de base|considerac"),
    ("Entregables e informes", r"informe|entrega|producto|avance|recepcion|exposici"),
    ("Multas y término",       r"multa|sancion|termino anticipado|incumplimiento|caducidad"),
    ("Equipo de trabajo",      r"equipo|profesional|personal clave|recurso humano"),
]
STOP_TITULOS = re.compile(r"^(puntos?|republica de chile|pagina)\b", re.I)

def sin_tildes(s):
    return unicodedata.normalize("NFD", s or "").encode("ascii", "ignore").decode()

def limpia_titulo(t):
    return re.sub(r"\s+", " ", t).strip(" .:-")

def es_titulo(t):
    """Heurística: título de sección = mayúsculas, con sustancia."""
    t = t.strip()
    if len(t) < 6 or len(t) > 85: return False
    if STOP_TITULOS.search(t): return False
    letras = [c for c in t if c.isalpha()]
    if not letras: return False
    up = sum(1 for c in letras if c.isupper())
    return up >= len(letras) * 0.7            # >=70% mayúsculas

_MARCA = ("-", "–", "—", "•", "▪", "◦", "*", "·", "o")
_ITEM = re.compile(r"^(•|\d{1,2}(?:\.\d{1,2}){0,2}[\.\)]\s|[a-zA-Z][\.\)]\s|"
                   r"TAREA\b|ETAPA\b|ANEXO\b|FORMULARIO\b|ART[ÍI]CULO\b|▪)", re.I)

def reflow(t):
    """Reflujo: junta fragmentos en párrafos y viñetas reales, sin huecos."""
    raw = [l.strip() for l in (t or "").split("\n")]
    lines = []; i = 0
    while i < len(raw):                          # 1) marcador solitario + línea siguiente
        l = raw[i]
        if l in _MARCA and i + 1 < len(raw) and raw[i + 1]:
            lines.append("• " + raw[i + 1]); i += 2; continue
        m = re.match(r"^[-–—•▪◦*·]\s+(.*)", l)
        if m:
            lines.append("• " + m.group(1)); i += 1; continue
        lines.append(l); i += 1
    out = []; buf = []                            # 2) acumula corridas en párrafos
    def flush():
        if buf:
            out.append(" ".join(buf)); buf.clear()
    for l in lines:
        if not l:
            flush(); out.append("")
        elif _ITEM.match(l):
            flush(); out.append(l)
        else:
            buf.append(l)
    flush()
    res = []                                      # 3) colapsa blancos múltiples
    for l in out:
        if l == "" and (not res or res[-1] == ""):
            continue
        res.append(l)
    return "\n".join(res).strip()

def secciones_de_texto(txt):
    """Detecta secciones numeradas; devuelve [{num, titulo, texto}]."""
    pat = re.compile(r"(?m)^\s*(\d{1,2}(?:\.\d{1,2}){0,2})[\.\)]?\s+"
                     r"([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ0-9 ,.:&\-()/°ª\"']{3,90})")
    cands = []
    for m in pat.finditer(txt):
        titulo = re.sub(r"\s*\d{1,3}\s*$", "", limpia_titulo(m.group(2)))  # quita nº de página final
        if es_titulo(titulo):
            cands.append((m.start(), m.end(), m.group(1), titulo))
    from collections import Counter
    frec = Counter(c[3] for c in cands)
    cands = [c for c in cands if frec[c[3]] <= 4]
    secs = []
    for i, (ini, fin_head, num, tit) in enumerate(cands):
        fin = cands[i + 1][0] if i + 1 < len(cands) else len(txt)
        secs.append({"num": num, "titulo": tit, "texto": reflow(txt[fin_head:fin])})
    return secs

def leer_pdf(path):
    import fitz
    doc = fitz.open(path)
    paginas = [pg.get_text() for pg in doc]
    doc.close()
    return "".join(paginas), paginas

def ocr_pdf(path):
    """OCR página a página (español) para PDFs escaneados."""
    import fitz
    doc = fitz.open(path); paginas = []
    n = doc.page_count
    for i, pg in enumerate(doc):
        try:
            tp = pg.get_textpage_ocr(language="spa", dpi=200, full=True)
            paginas.append(pg.get_text(textpage=tp))
        except Exception:
            paginas.append("")
        print(f"       OCR página {i + 1}/{n}", end="\r")
    doc.close(); print()
    return "".join(paginas), paginas

def leer_docx(path):
    import docx
    d = docx.Document(str(path))
    partes = [p.text for p in d.paragraphs]
    for tbl in d.tables:                       # los anexos suelen ser tablas
        for row in tbl.rows:
            partes.append(" | ".join(c.text for c in row.cells))
    return "\n".join(partes), None

def _safe(s):
    return re.sub(r'[\\/:*?"<>|]+', "_", s)[:120]

def _md_path(carpeta, archivo):
    return MD_DIR / _safe(carpeta) / (_safe(archivo) + ".md")

def guardar_md(carpeta, path, doc):
    """Guarda la base leída como Markdown (caché + lectura humana)."""
    p = _md_path(carpeta, path.name); p.parent.mkdir(parents=True, exist_ok=True)
    L = [f"<!-- radar-bases cls={doc['cls']} paginas={doc.get('paginas')} "
         f"ocr={int(bool(doc.get('ocr')))} escaneado={int(bool(doc.get('escaneado')))} -->",
         f"# {doc['archivo']}", ""]
    if doc.get("escaneado"):
        L += ["> " + (doc.get("nota") or ""), ""]
    for s in doc.get("secciones", []):
        L.append("## " + ((s["num"] + " · " if s["num"] else "") + s["titulo"]))
        for ln in s["texto"].split("\n"):
            L.append("- " + ln[2:] if ln.startswith("• ") else ln)
        L.append("")
    p.write_text("\n".join(L), encoding="utf-8")

def cargar_md(carpeta, path):
    """Devuelve (meta, secciones) si hay caché .md más nuevo que el original."""
    if REPROC:
        return None
    p = _md_path(carpeta, path.name)
    try:
        if not p.exists() or p.stat().st_mtime < path.stat().st_mtime:
            return None
    except Exception:
        return None
    txt = p.read_text(encoding="utf-8")
    m0 = re.match(r"<!-- radar-bases (.*?)-->", txt)
    meta = dict(re.findall(r"(\w+)=(\S+)", m0.group(1))) if m0 else {}
    secs = []
    for blk in re.split(r"(?m)^## ", txt)[1:]:
        head, _, body = blk.partition("\n")
        num, sep, tit = head.partition(" · ")
        if not sep:
            num, tit = "", num
        b = "\n".join("• " + l[2:] if l.startswith("- ") else l
                      for l in body.split("\n")).strip()
        secs.append({"num": num.strip(), "titulo": tit.strip(), "texto": b})
    return meta, secs

def clasificar_claves(secs):
    claves = {}
    for idx, s in enumerate(secs):
        low = sin_tildes(s["titulo"]).lower()
        for cat, pat in CATEGORIAS:
            if re.search(pat, low):
                snip = re.sub(r"\s+", " ", s["texto"])[:220]
                claves.setdefault(cat, []).append(
                    {"i": idx, "titulo": f'{s["num"]} {s["titulo"]}', "snippet": snip})
    return claves

def procesar_documento(path, carpeta):
    nombre, cls = tipo_doc(path.name)
    # 1) ¿hay caché markdown vigente? -> se lee de ahí, sin re-procesar
    cache = cargar_md(carpeta, path)
    if cache:
        meta, secs = cache
        pag = meta.get("paginas", "None")
        doc = {"archivo": path.name, "tipo": nombre, "cls": cls, "cache": True,
               "paginas": int(pag) if pag.isdigit() else None, "ocr": meta.get("ocr") == "1"}
        if meta.get("escaneado") == "1":
            doc.update(escaneado=True, secciones=[], claves={},
                       nota="Escaneado sin texto legible (revisa el original).")
        else:
            doc.update(secciones=secs, claves=clasificar_claves(secs))
        return doc
    # 2) sin caché -> leer/OCR y guardar markdown
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            txt, paginas = leer_pdf(path)
        elif ext == ".docx":
            txt, paginas = leer_docx(path)
        else:
            return None
    except Exception as e:
        print(f"     no se pudo leer {path.name}: {str(e)[:60]}")
        return None
    ocr_usado = False
    if len((txt or "").strip()) < 200 and ext == ".pdf" and OCR_OK:
        print(f"     {path.name[:40]}: escaneado -> OCR en español…")
        try:
            txt, paginas = ocr_pdf(path); ocr_usado = True
        except Exception as e:
            print(f"     OCR falló: {str(e)[:50]}")
    if len((txt or "").strip()) < 200:
        nota = ("Sin capa de texto y OCR no disponible (instala Tesseract)."
                if not OCR_OK else "No se pudo extraer texto ni con OCR.")
        doc = {"archivo": path.name, "tipo": nombre, "cls": cls,
               "paginas": len(paginas) if paginas else None,
               "escaneado": True, "secciones": [], "claves": {}, "nota": nota}
        guardar_md(carpeta, path, doc)
        return doc
    secs = secciones_de_texto(txt)
    if len(secs) < 3 and paginas:
        secs = [{"num": str(i + 1), "titulo": f"Página {i + 1}", "texto": reflow(p)}
                for i, p in enumerate(paginas) if p.strip()]
    if not secs:
        secs = [{"num": "", "titulo": "Contenido", "texto": reflow(txt)}]
    doc = {"archivo": path.name, "tipo": nombre, "cls": cls,
           "paginas": len(paginas) if paginas else None, "ocr": ocr_usado,
           "secciones": secs, "claves": clasificar_claves(secs)}
    guardar_md(carpeta, path, doc)
    return doc

def main():
    if not BASE_DIR.exists():
        sys.exit(f"No existe la carpeta {BASE_DIR}")
    print(f"Carpeta: {BASE_DIR}\n")
    licitaciones = []
    carpetas = [d for d in BASE_DIR.iterdir() if d.is_dir() and not d.name.startswith("_")]
    for carpeta in sorted(carpetas):
        # descomprime ZIPs si hay
        for z in carpeta.glob("*.zip"):
            try:
                with zipfile.ZipFile(z) as zf:
                    zf.extractall(carpeta / (z.stem + "_extraido"))
                print(f"  [zip] {z.name} descomprimido")
            except Exception as e:
                print(f"  [zip] error {z.name}: {e}")
        docs = []
        vistos = set()
        for f in sorted(carpeta.rglob("*")):
            if f.suffix.lower() in (".pdf", ".docx") and f.is_file():
                clave = f.stat().st_size            # dedup copias "(1)" idénticas
                if clave in vistos:
                    print(f"  (dup) se omite {f.name}"); continue
                vistos.add(clave)
                d = procesar_documento(f, carpeta.name)
                if d:
                    docs.append(d)
                    tag = "cache" if d.get("cache") else ("OCR" if d.get("ocr") else "leído")
                    nsec = len(d["secciones"]); nclv = sum(len(v) for v in d["claves"].values())
                    print(f"  [{tag:5}] {carpeta.name[:22]} / {f.name[:38]}  {nsec} secc, {nclv} claves")
        if docs:
            licitaciones.append({"carpeta": carpeta.name,
                                 "titulo": re.sub(r"^\d+\s*", "", carpeta.name),
                                 "documentos": docs})

    data = {"generado": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
            "licitaciones": licitaciones}
    salida_js = BASE_DIR / "bases_data.js"
    with open(salida_js, "w", encoding="utf-8") as fp:
        fp.write("window.BASES = ")
        json.dump(data, fp, ensure_ascii=False)
        fp.write(";")
    # copia el visor
    visor = AQUI / "analizador_bases.html"
    destino = None
    if visor.exists():
        destino = BASE_DIR / "Analizador de Bases.html"
        destino.write_text(visor.read_text(encoding="utf-8"), encoding="utf-8")
        print(f"\nVisor copiado a: {destino}")
    tot_doc = sum(len(l["documentos"]) for l in licitaciones)
    print(f"\nListo: {len(licitaciones)} licitaciones, {tot_doc} documentos.")
    print(f"Datos: {salida_js}")
    print(f"Markdown (caché legible): {MD_DIR}")
    if destino and "--no-abrir" not in sys.argv:
        try:
            os.startfile(str(destino))          # abre el visor (Windows)
        except Exception:
            print("Abre 'Analizador de Bases.html' en la carpeta de licitaciones.")

if __name__ == "__main__":
    main()
